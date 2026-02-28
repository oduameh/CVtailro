"""Professional DOCX resume generator.

Converts the same markdown format used by the PDF generator into a
professionally formatted Word document using python-docx.

Markdown format handled:
    # Name
    Location | Phone | Email | LinkedIn

    ## Professional Summary
    Paragraph text...

    ## Experience
    **Job Title** | Company | Location
    Start Date - End Date
    - Bullet point 1
    - Bullet point 2

    ## Education
    **Degree** -- School | Date

    ## Skills
    **Category:** Skill1, Skill2, Skill3
"""

from __future__ import annotations

import logging
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Helper: add a bottom border to a paragraph (used for section headings)
# ---------------------------------------------------------------------------

def _add_bottom_border(paragraph):
    """Add a thin bottom border to a paragraph via Open XML."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pBdr.makeelement(
        qn("w:bottom"),
        {
            qn("w:val"): "single",
            qn("w:sz"): "4",  # 0.5pt
            qn("w:space"): "1",
            qn("w:color"): "888888",
        },
    )
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set exact spacing on a paragraph (values in Pt)."""
    pPr = paragraph._p.get_or_add_pPr()
    spacing = pPr.makeelement(qn("w:spacing"), {})
    spacing.set(qn("w:before"), str(int(before * 20)))  # twips
    spacing.set(qn("w:after"), str(int(after * 20)))
    if line is not None:
        spacing.set(qn("w:line"), str(int(line * 20)))
        spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)


def _strip_bold_markers(text: str) -> str:
    """Remove ** bold markers from markdown text."""
    return re.sub(r"\*\*(.+?)\*\*", r"\1", text)


def _strip_inline_md(text: str) -> str:
    """Remove all inline markdown: bold, italic, links."""
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"\*(.+?)\*", r"\1", text)
    text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)
    return text


# ---------------------------------------------------------------------------
# Contact block parser (mirrors pdf_generator._parse_contact_block)
# ---------------------------------------------------------------------------

def _parse_contact_block(lines: list[str]) -> tuple[str, list[str], int]:
    """Extract name and contact info from the top of the resume."""
    name = ""
    contact_lines = []
    first_section_idx = 0
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            continue
        if stripped in ("---", "***", "___"):
            continue
        if re.match(r"^#{2,}\s+", stripped):
            first_section_idx = i
            break
        if stripped.startswith("#"):
            if not name:
                name = re.sub(r"^#+\s*", "", stripped)
                continue
            else:
                first_section_idx = i
                break
        if not name:
            name = stripped
        else:
            contact_lines.append(stripped)
    if first_section_idx == 0:
        first_section_idx = len(lines)
    return name, contact_lines, first_section_idx


def _flatten_contact(contact_lines: list[str]) -> str:
    """Flatten contact lines into a single pipe-separated string."""
    items = []
    for line in contact_lines:
        line = _strip_inline_md(line)
        if "|" in line:
            for part in line.split("|"):
                cleaned = part.strip()
                if cleaned:
                    items.append(cleaned)
        else:
            cleaned = line.strip()
            if cleaned:
                items.append(cleaned)
    return "  |  ".join(items)


# ---------------------------------------------------------------------------
# Main generator
# ---------------------------------------------------------------------------

def generate_resume_docx(md_content: str, output_path: str | Path) -> Path:
    """Generate a professionally formatted DOCX from markdown resume content.

    Args:
        md_content: Resume content in markdown format.
        output_path: Path for the output DOCX file.

    Returns:
        Path to the generated DOCX file.
    """
    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # --- Page margins: 0.6 inches all sides ---
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    # --- Default font ---
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    paragraph_format = style.paragraph_format
    paragraph_format.space_before = Pt(0)
    paragraph_format.space_after = Pt(2)

    lines = md_content.split("\n")
    name, contact_lines, start_idx = _parse_contact_block(lines)

    # --- Name ---
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_name.add_run(name)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    _set_paragraph_spacing(p_name, before=0, after=2)

    # --- Contact ---
    if contact_lines:
        contact_text = _flatten_contact(contact_lines)
        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p_contact.add_run(contact_text)
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        _set_paragraph_spacing(p_contact, before=0, after=6)

    # --- Parse body ---
    current_section = ""
    i = start_idx

    while i < len(lines):
        line = lines[i].rstrip()
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Horizontal rule -- skip
        if stripped in ("---", "***", "___"):
            i += 1
            continue

        # Section heading: ## Heading (not ### which is a role header)
        section_match = re.match(r"^#{1,2}\s+(.+)", line)
        if section_match and not re.match(r"^###\s+", line):
            heading_text = section_match.group(1).strip()
            current_section = heading_text.lower()
            p = doc.add_paragraph()
            run = p.add_run(heading_text.upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            _add_bottom_border(p)
            _set_paragraph_spacing(p, before=8, after=4)
            i += 1
            continue

        # Role header: ### Title | Company | Location | Date
        role_h3 = re.match(r"^###\s+(.+)", line)
        if role_h3 and (
            "experience" in current_section
            or "employment" in current_section
            or "work history" in current_section
        ):
            parts = [
                p.strip()
                for p in re.split(r"\s*\|\s*", re.sub(r"^###\s*", "", line))
                if p.strip()
            ]
            title = parts[0] if len(parts) > 0 else ""
            company = parts[1] if len(parts) > 1 else ""
            location = parts[2] if len(parts) > 2 else ""
            date = parts[3] if len(parts) > 3 else ""
            if not date and i + 1 < len(lines):
                nl = lines[i + 1].strip()
                if nl and not nl.startswith(("-", "*", "#")):
                    date = nl
                    i += 1

            # Role line with date right-aligned using a tab stop
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=4, after=1)
            # Set a right-aligned tab stop at the right margin
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(
                Inches(7.3 - 1.2),  # page width minus margins
                alignment=WD_TAB_ALIGNMENT.RIGHT,
            )
            run_title = p.add_run(_strip_bold_markers(title))
            run_title.bold = True
            run_title.font.size = Pt(11)
            run_title.font.name = "Calibri"
            run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            if date:
                p.add_run("\t")
                run_date = p.add_run(date)
                run_date.font.size = Pt(10)
                run_date.font.name = "Calibri"
                run_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            # Company line
            if company:
                company_loc = company
                if location:
                    company_loc += f", {location}"
                p2 = doc.add_paragraph()
                _set_paragraph_spacing(p2, before=0, after=2)
                run_comp = p2.add_run(company_loc)
                run_comp.font.size = Pt(11)
                run_comp.font.name = "Calibri"
                run_comp.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            i += 1
            continue

        # Role header: **Title** | Company | Location (ATS format)
        role3 = re.match(r"^\*\*(.+?)\*\*\s*\|\s*(.+?)\s*\|\s*(.+)", line)
        if role3 and (
            "experience" in current_section
            or "employment" in current_section
            or "work history" in current_section
        ):
            title, company, location = (
                role3.group(1),
                role3.group(2),
                role3.group(3),
            )
            date = ""
            if i + 1 < len(lines):
                nl = lines[i + 1].strip()
                if nl and not nl.startswith(("-", "*", "#", "**")):
                    date = nl
                    i += 1

            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=4, after=1)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(
                Inches(7.3 - 1.2),
                alignment=WD_TAB_ALIGNMENT.RIGHT,
            )
            run_title = p.add_run(title)
            run_title.bold = True
            run_title.font.size = Pt(11)
            run_title.font.name = "Calibri"
            run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            if date:
                p.add_run("\t")
                run_date = p.add_run(date)
                run_date.font.size = Pt(10)
                run_date.font.name = "Calibri"
                run_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            company_loc = company
            if location:
                company_loc += f", {location}"
            p2 = doc.add_paragraph()
            _set_paragraph_spacing(p2, before=0, after=2)
            run_comp = p2.add_run(company_loc)
            run_comp.font.size = Pt(11)
            run_comp.font.name = "Calibri"
            run_comp.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            i += 1
            continue

        # Role header: **Title** | Company (2-part)
        role2 = re.match(r"^\*\*(.+?)\*\*\s*\|\s*(.+)", line)
        if role2 and (
            "experience" in current_section
            or "employment" in current_section
            or "work history" in current_section
        ):
            title, company = role2.group(1), role2.group(2)
            date = ""
            if i + 1 < len(lines):
                nl = lines[i + 1].strip()
                if nl and not nl.startswith(("-", "*", "#", "**")):
                    date = nl
                    i += 1

            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=4, after=1)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(
                Inches(7.3 - 1.2),
                alignment=WD_TAB_ALIGNMENT.RIGHT,
            )
            run_title = p.add_run(title)
            run_title.bold = True
            run_title.font.size = Pt(11)
            run_title.font.name = "Calibri"
            run_title.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            if date:
                p.add_run("\t")
                run_date = p.add_run(date)
                run_date.font.size = Pt(10)
                run_date.font.name = "Calibri"
                run_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            p2 = doc.add_paragraph()
            _set_paragraph_spacing(p2, before=0, after=2)
            run_comp = p2.add_run(company)
            run_comp.font.size = Pt(11)
            run_comp.font.name = "Calibri"
            run_comp.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            i += 1
            continue

        # Education: **Degree** | School | Date  OR  **Degree** -- School | Date
        edu_match = re.match(r"^\*\*(.+?)\*\*\s*[|—\u2013-]+\s*(.+)", stripped)
        if edu_match and "education" in current_section:
            degree = edu_match.group(1)
            rest = edu_match.group(2)
            rest_parts = [
                p.strip()
                for p in re.split(r"\s*[|—\u2013]\s*", rest)
                if p.strip()
            ]
            school = rest_parts[0] if rest_parts else ""
            date = rest_parts[1] if len(rest_parts) > 1 else ""
            if not date and i + 1 < len(lines):
                nl = lines[i + 1].strip()
                if nl and not nl.startswith(("-", "*", "#", "**")):
                    date = nl
                    i += 1

            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=4, after=1)
            tab_stops = p.paragraph_format.tab_stops
            tab_stops.add_tab_stop(
                Inches(7.3 - 1.2),
                alignment=WD_TAB_ALIGNMENT.RIGHT,
            )
            run_deg = p.add_run(degree)
            run_deg.bold = True
            run_deg.font.size = Pt(11)
            run_deg.font.name = "Calibri"
            run_deg.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            if date:
                p.add_run("\t")
                run_date = p.add_run(date)
                run_date.font.size = Pt(10)
                run_date.font.name = "Calibri"
                run_date.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

            if school:
                p2 = doc.add_paragraph()
                _set_paragraph_spacing(p2, before=0, after=2)
                run_sch = p2.add_run(school)
                run_sch.font.size = Pt(10)
                run_sch.font.name = "Calibri"
                run_sch.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

            i += 1
            continue

        # Certification: **Name** | meta  OR  **Name** -- meta
        if "certification" in current_section and stripped:
            cert_bold = re.match(r"^\*\*(.+?)\*\*\s*[|—\u2013-]+\s*(.+)", stripped)
            cert_plain = (
                re.match(r"^(.+?)\s*[—\u2013]\s*(.+)", stripped)
                if not cert_bold
                else None
            )
            if cert_bold:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=2, after=2)
                run_name = p.add_run(cert_bold.group(1))
                run_name.bold = True
                run_name.font.size = Pt(10)
                run_name.font.name = "Calibri"
                p.add_run(" -- ")
                run_meta = p.add_run(cert_bold.group(2))
                run_meta.font.size = Pt(10)
                run_meta.font.name = "Calibri"
                run_meta.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            elif cert_plain:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=2, after=2)
                run_name = p.add_run(cert_plain.group(1))
                run_name.bold = True
                run_name.font.size = Pt(10)
                run_name.font.name = "Calibri"
                p.add_run(" -- ")
                run_meta = p.add_run(cert_plain.group(2))
                run_meta.font.size = Pt(10)
                run_meta.font.name = "Calibri"
                run_meta.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                p = doc.add_paragraph()
                _set_paragraph_spacing(p, before=2, after=2)
                run_name = p.add_run(_strip_inline_md(stripped))
                run_name.bold = True
                run_name.font.size = Pt(10)
                run_name.font.name = "Calibri"
            i += 1
            continue

        # Bullet point
        bullet_match = re.match(r"^\s*[-*]\s+(.+)", line)
        if bullet_match:
            bullet_text = _strip_inline_md(bullet_match.group(1))
            p = doc.add_paragraph(style="List Bullet")
            _set_paragraph_spacing(p, before=0, after=1)
            # Clear default run and add our own with proper formatting
            p.clear()
            run = p.add_run(bullet_text)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            # Set left indent for proper bullet indentation
            p.paragraph_format.left_indent = Inches(0.25)
            i += 1
            continue

        # Skills -- categorized: **Category:** items
        cat_skill = re.match(r"^\*\*(.+?)\*\*\s*:\s*(.+)", stripped)
        if cat_skill and "skill" in current_section:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=1, after=2)
            run_cat = p.add_run(cat_skill.group(1) + ": ")
            run_cat.bold = True
            run_cat.font.size = Pt(10)
            run_cat.font.name = "Calibri"
            run_cat.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
            run_items = p.add_run(cat_skill.group(2).strip())
            run_items.font.size = Pt(10)
            run_items.font.name = "Calibri"
            run_items.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            i += 1
            continue

        # Skills -- plain comma-separated
        if (
            stripped
            and "skill" in current_section
            and not re.match(r"^\*\*", stripped)
        ):
            skill_text = stripped
            while i + 1 < len(lines):
                nl = lines[i + 1].strip()
                if not nl or nl.startswith(("#", "**")):
                    break
                i += 1
                skill_text += " " + nl
            cleaned = re.sub(r"\s*,\s*", ", ", skill_text.strip().rstrip(","))
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=1, after=2)
            run = p.add_run(cleaned)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
            i += 1
            continue

        # Skills -- bold category fallback
        if (
            stripped
            and "skill" in current_section
            and re.match(r"^\*\*", stripped)
        ):
            clean = _strip_inline_md(stripped)
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=1, after=2)
            run = p.add_run(clean)
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            i += 1
            continue

        # Summary -- consolidate consecutive lines
        if stripped and (
            "summary" in current_section
            or "profile" in current_section
            or "objective" in current_section
        ):
            summary_text = stripped
            while i + 1 < len(lines):
                nl = lines[i + 1].strip()
                if not nl or nl.startswith(("#", "---", "***", "___")):
                    break
                i += 1
                summary_text += " " + nl
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=4)
            run = p.add_run(_strip_inline_md(summary_text))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            i += 1
            continue

        # Generic text
        if stripped:
            p = doc.add_paragraph()
            _set_paragraph_spacing(p, before=0, after=2)
            run = p.add_run(_strip_inline_md(stripped))
            run.font.size = Pt(10)
            run.font.name = "Calibri"
            run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

        i += 1

    # Save
    doc.save(str(out))

    if not out.exists() or out.stat().st_size < 100:
        raise RuntimeError(f"DOCX generation produced empty or missing file: {out}")

    logger.info(f"DOCX generated: {out} ({out.stat().st_size} bytes)")
    return out
